import os
import re
import docx
import shutil
import zipfile

import sys
import platform
import subprocess

# --- CONSTANTS & PATHS ---
if getattr(sys, 'frozen', False):
    # Running as executable (Frozen)
    ext_root = os.path.dirname(sys.executable)
    portable_data_dir = os.path.join(ext_root, "data")
    internal_dir = os.path.join(ext_root, "_internal")
    
    # Locate Bundled Templates (Source)
    if hasattr(sys, '_MEIPASS'):
        # PyInstaller --onefile mode
        BUNDLED_TEMPLATES_DIR = os.path.join(sys._MEIPASS, "templates")
        BUNDLED_CITATION_DIR = os.path.join(sys._MEIPASS, "citation")
    else:
        # PyInstaller --onedir mode
        # Check _internal/templates (Windows newer default) or adjacent templates
        check_bundled = os.path.join(internal_dir, "templates")
        if not os.path.exists(check_bundled):
             # Try adjacent to executable (macOS often places here with --add-data)
             check_bundled = os.path.join(ext_root, "templates")
        BUNDLED_TEMPLATES_DIR = check_bundled
        
        check_bundled_cit = os.path.join(internal_dir, "citation")
        if not os.path.exists(check_bundled_cit):
             check_bundled_cit = os.path.join(ext_root, "citation")
        BUNDLED_CITATION_DIR = check_bundled_cit

    # Determine Runtime Storage (Destination)
    # Priority 1: 'data' folder next to exe (Portable Mode - Explicit)
    if os.path.exists(portable_data_dir):
        BASE_DIR = portable_data_dir
    # Priority 2: '_internal/templates' exists (Portable Mode - Implicit/Default)
    elif os.path.exists(os.path.join(internal_dir, "templates")):
        BASE_DIR = internal_dir
    else:
        # Priority 3: Installed/Standard Mode (Fallback)
        if platform.system() == 'Windows':
            base_storage = os.environ.get('APPDATA', os.path.expanduser("~")) 
            BASE_DIR = os.path.join(base_storage, "LegalDocGen")
        elif platform.system() == 'Darwin':
            # macOS: Use Documents to avoid read-only errors and ensure user accessibility
            BASE_DIR = os.path.expanduser("~/Documents/LegalDocAutomator")
        else:
            base_storage = os.path.expanduser("~/.config")
            BASE_DIR = os.path.join(base_storage, "LegalDocGen")
            
else:
    # Running as script (Dev)
    BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    BUNDLED_TEMPLATES_DIR = os.path.join(BASE_DIR, "templates")
    BUNDLED_CITATION_DIR = os.path.join(BASE_DIR, "citation")

CONFIG_DIR = os.path.join(BASE_DIR, "config")
TEMPLATE_DIR = os.path.join(BASE_DIR, "templates")
CITATION_DIR = os.path.join(BASE_DIR, "citation")

def ensure_directories():
    for d in [CONFIG_DIR, TEMPLATE_DIR, CITATION_DIR]:
        os.makedirs(d, exist_ok=True)
    
    # Auto-initialize templates if empty (First Run)
    if os.path.exists(BUNDLED_TEMPLATES_DIR) and os.path.exists(TEMPLATE_DIR):
        if not os.listdir(TEMPLATE_DIR):
            try:
                print(f"First run detected. Copying templates from {BUNDLED_TEMPLATES_DIR} to {TEMPLATE_DIR}")
                # We iterate and copy to avoid deleting the destination root
                for item in os.listdir(BUNDLED_TEMPLATES_DIR):
                    s = os.path.join(BUNDLED_TEMPLATES_DIR, item)
                    d = os.path.join(TEMPLATE_DIR, item)
                    if os.path.isdir(s):
                        shutil.copytree(s, d, dirs_exist_ok=True)
                    else:
                        shutil.copy2(s, d)
            except Exception as e:
                print(f"Error copying default templates: {e}")
                
    # Auto-initialize citation library if empty (First Run)
    if os.path.exists(BUNDLED_CITATION_DIR) and os.path.exists(CITATION_DIR):
        if not os.listdir(CITATION_DIR):
            try:
                print(f"First run detected. Copying citation library from {BUNDLED_CITATION_DIR} to {CITATION_DIR}")
                for item in os.listdir(BUNDLED_CITATION_DIR):
                    s = os.path.join(BUNDLED_CITATION_DIR, item)
                    d = os.path.join(CITATION_DIR, item)
                    if os.path.isdir(s):
                        shutil.copytree(s, d, dirs_exist_ok=True)
                    else:
                        shutil.copy2(s, d)
            except Exception as e:
                print(f"Error copying default citation library: {e}")

class SmartSync:
    @staticmethod
    def sync_template_variable(template_path, old_var, new_var):
        """
        Surgically replaces {{old_var}} with {{new_var}} inside a docx file
        while preserving run formatting (bold/color/etc).
        """
        if not os.path.exists(template_path):
            return False
            
        ext = os.path.splitext(template_path)[1].lower()
        if ext != '.docx':
            return False # Only supported for docx right now

        try:
            doc = docx.Document(template_path)
        except Exception:
            return False

        modified = False

        # 1. Paragraphs
        for p in doc.paragraphs:
            if SmartSync._replace_in_paragraph(p, old_var, new_var):
                modified = True

        # 2. Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if SmartSync._replace_in_paragraph(p, old_var, new_var):
                            modified = True
                            
        if modified:
            try:
                doc.save(template_path)
                return True
            except Exception:
                pass
        return False

    @staticmethod
    def _replace_in_paragraph(paragraph, old_var, new_var):
        """
        Helper: Replaces text in a single paragraph object across multiple runs.
        """
        full_text = paragraph.text
        escaped_var = re.escape(old_var)
        pattern = rf"\{{\{{\s*{escaped_var}\s*\}}\}}"
        
        if not re.search(pattern, full_text):
            return False

        replacement = f"{{{{{new_var}}}}}"
        modified = False
        
        # Simple case: if the tag is fully inside a run
        for run in paragraph.runs:
            if re.search(pattern, run.text):
                run.text = re.sub(pattern, replacement, run.text)
                modified = True

        # If it wasn't modified in runs, or there are still occurrences wrapping across runs
        new_full_text = paragraph.text
        if re.search(pattern, new_full_text):
            # Fallback: naive full-paragraph replace if split across runs.
            paragraph.text = re.sub(pattern, replacement, new_full_text)
            modified = True
                
        return modified

def open_file_or_folder(path):
    """
    Open a file or folder using the default system application in a cross-platform way.
    """
    if platform.system() == 'Windows':
        os.startfile(path)
    elif platform.system() == 'Darwin':       # macOS
        subprocess.call(['open', path])
    else:                                     # Linux
        subprocess.call(['xdg-open', path])

def extract_fields(path):
    """
    Parses a docx or odt file and returns a set of all {{variable}} names found.
    """
    if not os.path.exists(path):
        return set()
    
    unique_fields = set()
    ext = os.path.splitext(path)[1].lower()

    if ext == '.docx':
        return _extract_from_docx(path)
    elif ext == '.odt':
        return _extract_from_odt(path)
    return set()

def _extract_from_odt(path):
    unique_fields = set()
    try:
        with zipfile.ZipFile(path, 'r') as z:
            content = z.read('content.xml').decode('utf-8')
            # Regex for {{...}}
            matches = re.findall(r"\{\{(.*?)\}\}", content)
            for m in matches:
                clean = m.strip()
                # ODT might insert XML tags inside the braces if formatted?
                # Simpler regex might strip basic XML tags if they appear, 
                # but usually py3o expects clean jinja tags.
                # Let's assume user inputs clean tags for now or use a robust cleaner.
                # However, XML tags inside {{ }} break Jinja.
                # Users should type cleanly.
                if clean:
                    unique_fields.add(clean)
    except Exception as e:
        print(f"Error parsing ODT {path}: {e}")
    return unique_fields

def _extract_from_docx(path):
    """
    Parses a docx file and returns a set of all {{variable}} names found.
    """
    if not os.path.exists(path):
        return set()
    
    unique_fields = set()
    try:
        # Use a context manager to ensure the file handle is closed.
        # python-docx loads the file into memory, so we can close the handle after loading (or after processing).
        with open(path, 'rb') as f:
            doc = docx.Document(f)
            
            # Helper to scan text
            def scan_text(text):
                # Find all {{...}} patterns
                matches = re.findall(r"\{\{(.*?)\}\}", text)
                for m in matches:
                    # m is the content inside brackets. 
                    # It might be " Name " -> strip it to "Name"
                    clean = m.strip()
                    if clean:
                        unique_fields.add(clean)

            # 1. Paragraphs
            for p in doc.paragraphs:
                scan_text(p.text)
                
            # 2. Tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            scan_text(p.text)
                        
    except Exception as e:
        print(f"Error parsing {path}: {e}")
        
    return unique_fields
